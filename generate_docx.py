import sys, requests, os, json, base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_numbered_heading(doc, text, level, num_str):
    p = doc.add_heading(level=level)
    run = p.add_run(f"{num_str} {text}")
    run.font.color.rgb = RGBColor(31, 56, 100)
    run.font.size = Pt(16) if level == 1 else Pt(13)

def safe_add_text(doc, data):
    """Ensures dictionaries or lists are converted to string before adding to paragraph"""
    if isinstance(data, (dict, list)):
        text = json.dumps(data, indent=2)
    else:
        text = str(data)
    doc.add_paragraph(text)

def build_docx(iflow_name, author, date, ai_data, output_path):
    doc = Document()
    
    # --- HEADER LOGOS ---
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, width=Inches(6.5))
    sap_url = "https://raw.githubusercontent.com/Akila710/CPI_Demo1/main/assets/logos/SAP.jpg"
    mm_url = "https://raw.githubusercontent.com/Akila710/CPI_Demo1/main/assets/logos/mm_logo.png"
    try:
        with open("sap.jpg", "wb") as f: f.write(requests.get(sap_url).content)
        with open("mm.png", "wb") as f: f.write(requests.get(mm_url).content)
        htable.rows[0].cells[0].paragraphs[0].add_run().add_picture("sap.jpg", height=Inches(0.75))
        mm_p = htable.rows[0].cells[1].paragraphs[0]
        mm_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        mm_p.add_run().add_picture("mm.png", height=Inches(0.75))
    except: pass

    # --- PAGE 1: COVER ---
    for _ in range(5): doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(iflow_name)
    run.font.size = Pt(32); run.font.bold = True; run.font.color.rgb = RGBColor(31, 56, 100)
    doc.add_paragraph("Technical Specification Document").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER; meta.style = 'Table Grid'
    for i, (l, v) in enumerate([("Author:", author), ("Date:", date), ("Version:", "1.0")]):
        meta.rows[i].cells[0].text = l; meta.rows[i].cells[1].text = v
    doc.add_page_break()

    # --- PAGE 2: TOC ---
    doc.add_heading("Table of Contents", level=1)
    toc = ["1. Introduction", "  1.1 Purpose", "  1.2 Scope", "2. Integration Overview", "  2.1 Integration Architecture", "  2.2 Integration Components", "3. Integration Scenarios", "  3.1 Scenario Description", "  3.2 Data Flows", "  3.3 Security Requirements", "4. Error Handling", "5. Testing Validation"]
    for item in toc:
        p = doc.add_paragraph(item)
        if item.startswith("  "): p.paragraph_format.left_indent = Inches(0.3)
    doc.add_page_break()

    # --- PAGE 3+: DYNAMIC SECTIONS ---
    add_numbered_heading(doc, "Introduction", 1, "1.")
    add_numbered_heading(doc, "Purpose", 2, "1.1")
    safe_add_text(doc, ai_data.get("purpose", ""))
    
    add_numbered_heading(doc, "Scope", 2, "1.2")
    safe_add_text(doc, ai_data.get("scope", ""))

    # 2. Integration Overview & DIAGRAM
    add_numbered_heading(doc, "Integration Overview", 1, "2.")
    add_numbered_heading(doc, "Integration Architecture", 2, "2.1")
    
    mermaid_code = ai_data.get("mermaid_diagram", "graph LR; Sender-->CPI; CPI-->Receiver;")
    graph_url = f"https://mermaid.ink/img/{base64.b64encode(mermaid_code.encode()).decode()}"
    try:
        with open("diagram.png", "wb") as f: f.write(requests.get(graph_url).content)
        doc.add_picture("diagram.png", width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except: doc.add_paragraph("[Diagram Generation Failed]")

    add_numbered_heading(doc, "Integration Components", 2, "2.2")
    # Create a table for components
    comp_data = ai_data.get("components", {})
    if isinstance(comp_data, dict):
        c_table = doc.add_table(rows=1, cols=2)
        c_table.style = 'Table Grid'
        c_table.rows[0].cells[0].text = 'Component'
        c_table.rows[0].cells[1].text = 'Value'
        for k, v in comp_data.items():
            row = c_table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
    else: safe_add_text(doc, comp_data)

    # 3. Integration Scenarios
    add_numbered_heading(doc, "Integration Scenarios", 1, "3.")
    add_numbered_heading(doc, "Scenario Description", 2, "3.1")
    safe_add_text(doc, ai_data.get("scenario_description", ""))
    
    add_numbered_heading(doc, "Data Flows", 2, "3.2")
    # Adding Mapping Table
    map_data = ai_data.get("mapping", [])
    if isinstance(map_data, list) and len(map_data) > 0:
        m_table = doc.add_table(rows=1, cols=3)
        m_table.style = 'Table Grid'
        m_table.rows[0].cells[0].text = 'Source Field'
        m_table.rows[0].cells[1].text = 'Target Field'
        m_table.rows[0].cells[2].text = 'Logic'
        for entry in map_data:
            row = m_table.add_row().cells
            row[0].text = str(entry.get("source", ""))
            row[1].text = str(entry.get("target", ""))
            row[2].text = str(entry.get("logic", ""))
    else: safe_add_text(doc, ai_data.get("data_flows", ""))

    add_numbered_heading(doc, "Security Requirements", 2, "3.3")
    safe_add_text(doc, ai_data.get("security", ""))

    # 4. Error Handling
    add_numbered_heading(doc, "Error Handling and Logging", 1, "4.")
    safe_add_text(doc, ai_data.get("error_handling", ""))

    # 5. Testing Validation
    add_numbered_heading(doc, "Testing Validation", 1, "5.")
    safe_add_text(doc, ai_data.get("testing", ""))

    doc.save(output_path)

if __name__ == "__main__":
    build_docx(sys.argv[1], sys.argv[2], sys.argv[3], json.loads(sys.argv[4]), sys.argv[5])
