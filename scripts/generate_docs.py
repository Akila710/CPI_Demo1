import os
import time
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

# ================= CONFIG =================
BASE_DIR = "cpi-artifacts"
SAP_LOGO = "assets/logos/SAP.jpg"
MM_LOGO = "assets/logos/mm_logo.png"

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
PACKAGE_NAME = os.getenv("PACKAGE_NAME")

TODAY = date.today().isoformat()
VERSION = "Draft"
AUTHOR = ""

# ================= VALIDATION =================
if not PACKAGE_NAME:
    raise Exception("PACKAGE_NAME not provided")

PACKAGE_PATH = os.path.join(BASE_DIR, PACKAGE_NAME)
if not os.path.isdir(PACKAGE_PATH):
    raise Exception(f"Package not found: {PACKAGE_PATH}")

iflows = [d for d in os.listdir(PACKAGE_PATH)
          if os.path.isdir(os.path.join(PACKAGE_PATH, d))]

if not iflows:
    raise Exception("No iFlows found")

print(f"Found iFlows: {iflows}")

# ================= GROQ CALL =================
def groq_generate(iflow_name, content):
    payload = {
        "model": "llama-3.3-70b-versatile",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a SAP CPI Technical Architect. "
                    "Generate clean plain-text documentation ONLY. "
                    "No markdown symbols, no bullets, no stars."
                )
            },
            {
                "role": "user",
                "content": f"""
Generate documentation strictly with numbered sections:

1. Introduction
1.1 Purpose
1.2 Scope

2. Integration Overview
2.1 Integration Architecture
2.2 Integration Components

3. Integration Scenarios
3.1 Scenario Description
3.2 Data Flows
3.3 Security Requirements

4. Error Handling and Logging
5. Testing Validation
6. Reference Documents

iFlow Name: {iflow_name}

iFlow Content:
{content}
"""
            }
        ]
    }

    r = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        },
        json=payload,
        timeout=60
    )

    if r.status_code == 429:
        print(f"⚠️ Groq rate limit hit for {iflow_name}, skipping")
        return None

    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

# ================= DOC HELPERS =================
def add_header(section):
    header = section.header
    header.paragraphs.clear()

    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(SAP_LOGO, width=Inches(1.1))

    #p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(MM_LOGO, width=Inches(1.1))

def add_toc_page(doc):
    p = doc.add_paragraph("Table of Contents")
    p.runs[0].bold = True

    toc = [
        "1. Introduction",
        "    1.1 Purpose",
        "    1.2 Scope",
        "2. Integration Overview",
        "    2.1 Integration Architecture",
        "    2.2 Integration Components",
        "3. Integration Scenarios",
        "    3.1 Scenario Description",
        "    3.2 Data Flows",
        "    3.3 Security Requirements",
        "4. Error Handling and Logging",
        "5. Testing Validation",
        "6. Reference Documents"
    ]

    for line in toc:
        doc.add_paragraph(line)

# ================= MAIN LOOP =================
for iflow in iflows:
    iflow_dir = os.path.join(PACKAGE_PATH, iflow)

    iflw_file = None
    for root, _, files in os.walk(iflow_dir):
        for f in files:
            if f.endswith(".iflw"):
                iflw_file = os.path.join(root, f)

    if not iflw_file:
        print(f"⚠️ No .iflw found for {iflow}, skipping")
        continue

    with open(iflw_file, encoding="utf-8") as f:
        iflw_content = f.read()

    content = groq_generate(iflow, iflw_content)
    if not content:
        continue

    # ---------- MD ----------
    md_path = os.path.join(iflow_dir, f"{iflow}.md")
    with open(md_path, "w", encoding="utf-8") as md:
        md.write(f"{iflow}\n\n{content}")

    # ---------- DOCX ----------
    doc = Document()
    section = doc.sections[0]
    add_header(section)

    # Cover page
    doc.add_paragraph("\n\n")
    title = doc.add_paragraph(iflow)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    title.runs[0].bold = True

    doc.add_paragraph("\n")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Author"
    table.cell(0, 1).text = AUTHOR
    table.cell(1, 0).text = "Date"
    table.cell(1, 1).text = TODAY
    table.cell(2, 0).text = "Version"
    table.cell(2, 1).text = VERSION

    doc.add_page_break()
    add_toc_page(doc)
    doc.add_page_break()

    for line in content.split("\n"):
        p = doc.add_paragraph(line)
        if line.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.")):
            p.runs[0].bold = True

    doc.save(os.path.join(iflow_dir, f"{iflow}.docx"))
    print(f"✅ Generated for {iflow}")

    time.sleep(2)  # ✅ avoid Groq rate limit

print("🎉 All possible iFlows processed")
